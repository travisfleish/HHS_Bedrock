#!/usr/bin/env python3
"""
OMHA Exhibit Tagger - OCR and Document Processing Test Script
Tests the enhanced document processing capabilities including OCR
"""

import os
import sys
import base64
import json
import requests
import time
from pathlib import Path
from typing import Dict, List, Tuple
import argparse

# Add the app directory to Python path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

try:
    from PIL import Image, ImageDraw, ImageFont
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    import PyPDF2
    from docx import Document
    import io
except ImportError as e:
    print(f"Missing test dependencies: {e}")
    print("Install with: pip install Pillow reportlab PyPDF2 python-docx")
    sys.exit(1)


class DocumentTestGenerator:
    """Generate test documents for OCR testing"""

    @staticmethod
    def create_text_pdf(text: str, filename: str) -> bytes:
        """Create a text-based PDF"""
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)

        # Add text
        y_position = 750
        for line in text.split('\n'):
            c.drawString(50, y_position, line)
            y_position -= 20

        c.save()
        buffer.seek(0)
        return buffer.read()

    @staticmethod
    def create_scanned_pdf(text: str, filename: str) -> bytes:
        """Create a scanned-looking PDF (image-based)"""
        # Create image with text
        img_width, img_height = 2550, 3300  # 8.5x11 at 300 DPI
        img = Image.new('RGB', (img_width, img_height), color='white')
        draw = ImageDraw.Draw(img)

        # Try to use a font, fall back to default if not available
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 60)
        except:
            font = ImageFont.load_default()

        # Add some noise to make it look scanned
        import random
        for _ in range(1000):
            x = random.randint(0, img_width)
            y = random.randint(0, img_height)
            gray = random.randint(200, 255)
            draw.point((x, y), fill=(gray, gray, gray))

        # Draw text
        y_position = 200
        for line in text.split('\n'):
            draw.text((200, y_position), line, fill='black', font=font)
            y_position += 100

        # Convert to PDF
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)

        # Save image to buffer
        img_buffer = io.BytesIO()
        img.resize((int(img_width / 3.5), int(img_height / 3.5)), Image.Resampling.LANCZOS).save(img_buffer,
                                                                                                 format='PNG')
        img_buffer.seek(0)

        # Add image to PDF
        c.drawImage(ImageReader(img_buffer), 0, 0, width=letter[0], height=letter[1])
        c.save()

        buffer.seek(0)
        return buffer.read()

    @staticmethod
    def create_mixed_pdf(text_pages: List[str], scanned_pages: List[str]) -> bytes:
        """Create a PDF with both text and scanned pages"""
        # This would require PyPDF2 to merge pages
        # For simplicity, we'll just create a text PDF with a note
        text = "MIXED PDF TEST\n\nPage 1 (Text-based):\n" + text_pages[0]
        text += "\n\nPage 2 (Would be scanned):\n" + scanned_pages[0] if scanned_pages else ""
        return DocumentTestGenerator.create_text_pdf(text, "mixed.pdf")

    @staticmethod
    def create_docx(text: str, filename: str) -> bytes:
        """Create a DOCX file"""
        doc = Document()
        doc.add_heading('Test Document', 0)

        for paragraph in text.split('\n\n'):
            doc.add_paragraph(paragraph)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()


class OCRTestSuite:
    """Test suite for OCR and document processing"""

    def __init__(self, base_url: str = "http://localhost:8001"):
        self.base_url = base_url
        self.test_results = []

    def run_health_check(self) -> Dict:
        """Check if the service is running and OCR is available"""
        try:
            response = requests.get(f"{self.base_url}/api/v1/health")
            response.raise_for_status()
            return response.json()
        except Exception as e:
            return {"error": str(e)}

    def test_document_extraction(self, content: bytes, mime_type: str, filename: str) -> Dict:
        """Test document extraction endpoint"""
        # Encode content to base64
        content_b64 = base64.b64encode(content).decode('utf-8')

        # Test using the debug endpoint if available
        try:
            response = requests.post(
                f"{self.base_url}/api/v1/test-extraction",
                json={
                    "file_content": content_b64,
                    "mime_type": mime_type
                }
            )
            if response.status_code == 200:
                return response.json()
        except:
            pass

        # Fall back to regular classification endpoint
        return self.test_classification(content_b64, mime_type, filename)

    def test_classification(self, content_b64: str, mime_type: str, filename: str) -> Dict:
        """Test full document classification"""
        request_data = {
            "case_id": "TEST-001",
            "documents": [{
                "filename": filename,
                "content": content_b64,
                "mime_type": mime_type
            }],
            "extract_detailed_data": True
        }

        try:
            start_time = time.time()
            response = requests.post(
                f"{self.base_url}/api/v1/tag-exhibits",
                json=request_data
            )
            elapsed_time = time.time() - start_time

            if response.status_code == 200:
                result = response.json()
                classification = result['classifications'][0] if result['classifications'] else {}

                return {
                    "success": True,
                    "filename": filename,
                    "processing_time": elapsed_time,
                    "classification": classification.get('document_type', 'unknown'),
                    "confidence": classification.get('confidence', 0),
                    "extraction_method": classification.get('extracted_data', {}).get('_extraction_method', 'unknown'),
                    "extraction_confidence": classification.get('extracted_data', {}).get('_extraction_confidence', 0),
                    "text_preview": classification.get('key_indicators', [])[:3]
                }
            else:
                return {
                    "success": False,
                    "filename": filename,
                    "error": f"HTTP {response.status_code}: {response.text}"
                }

        except Exception as e:
            return {
                "success": False,
                "filename": filename,
                "error": str(e)
            }

    def run_all_tests(self) -> None:
        """Run comprehensive test suite"""
        print("=" * 80)
        print("OMHA EXHIBIT TAGGER - OCR TEST SUITE")
        print("=" * 80)

        # 1. Health Check
        print("\n1. HEALTH CHECK")
        print("-" * 40)
        health = self.run_health_check()
        if "error" in health:
            print(f"❌ Service not available: {health['error']}")
            print("\nMake sure the service is running with: uvicorn app.main:app --reload")
            return

        print(f"✅ Service Status: {health.get('status', 'unknown')}")
        print(f"   OCR Available: {health.get('ocr_available', False)}")
        print(f"   OCR Enabled: {health.get('ocr_enabled', False)}")
        print(f"   Features: {json.dumps(health.get('features', {}), indent=6)}")

        if not health.get('ocr_available'):
            print("\n⚠️  WARNING: OCR is not available. Run setup_ocr.sh to install dependencies.")

        # 2. Generate Test Documents
        print("\n2. GENERATING TEST DOCUMENTS")
        print("-" * 40)

        # Sample Medicare text
        medicare_text = """MEDICARE RECONSIDERATION DECISION

Date: January 8, 2013
Beneficiary: JOHN DOE
HIC#: 123456789A
Appeal Number: 1-1221724503

This letter is to inform you of the decision on your Medicare Appeal. 
The appeal decision is UNFAVORABLE. Our decision is that Medicare will 
make no additional payment.

Reason for Denial:
The documentation does not support medical necessity for the High Frequency 
Chest Wall Oscillation Device (HFCWO) based on Local Coverage Determination 
(LCD) policy L12934.

You have the right to appeal this decision to an Administrative Law Judge (ALJ).
"""

        clinical_note_text = """CLINICAL PROGRESS NOTE

Patient: Jane Smith
DOB: 01/15/1945
Date of Service: 11/03/2023

Chief Complaint: Follow-up for bronchiectasis

History of Present Illness:
79-year-old female with history of bronchiectasis presents for routine follow-up.
Patient reports stable symptoms with occasional productive cough.

Assessment and Plan:
1. Bronchiectasis - stable
2. Continue current therapy
3. Follow-up in 6 months

Jorge Hernandez, MD
"""

        test_documents = []

        # Create text-based PDF
        print("   Creating text-based PDF...")
        text_pdf = DocumentTestGenerator.create_text_pdf(medicare_text, "medicare_text.pdf")
        test_documents.append(("medicare_text.pdf", text_pdf, "application/pdf", "text-based"))

        # Create scanned PDF
        print("   Creating scanned PDF...")
        scanned_pdf = DocumentTestGenerator.create_scanned_pdf(medicare_text, "medicare_scanned.pdf")
        test_documents.append(("medicare_scanned.pdf", scanned_pdf, "application/pdf", "scanned"))

        # Create DOCX
        print("   Creating DOCX file...")
        docx_file = DocumentTestGenerator.create_docx(clinical_note_text, "clinical_note.docx")
        test_documents.append(("clinical_note.docx", docx_file,
                               "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"))

        # Create plain text
        print("   Creating plain text file...")
        text_file = medicare_text.encode('utf-8')
        test_documents.append(("medicare.txt", text_file, "text/plain", "text"))

        # 3. Test Each Document
        print("\n3. TESTING DOCUMENT PROCESSING")
        print("-" * 40)

        for filename, content, mime_type, expected_type in test_documents:
            print(f"\nTesting: {filename}")
            print(f"   Expected type: {expected_type}")
            print(f"   Size: {len(content)} bytes")

            result = self.test_classification(
                base64.b64encode(content).decode('utf-8'),
                mime_type,
                filename
            )

            if result['success']:
                print(f"   ✅ Success!")
                print(f"   Classification: {result['classification']}")
                print(f"   Confidence: {result['confidence']:.2f}")
                print(f"   Extraction Method: {result['extraction_method']}")
                print(f"   Extraction Confidence: {result['extraction_confidence']:.2f}")
                print(f"   Processing Time: {result['processing_time']:.2f}s")
                if result['text_preview']:
                    print(f"   Key Indicators: {', '.join(result['text_preview'][:3])}")

                # Check if OCR was used for scanned PDF
                if expected_type == "scanned" and result['extraction_method'] != 'ocr':
                    print(f"   ⚠️  WARNING: Expected OCR extraction but got {result['extraction_method']}")
            else:
                print(f"   ❌ Failed: {result['error']}")

            self.test_results.append(result)

        # 4. Summary
        print("\n4. TEST SUMMARY")
        print("-" * 40)
        successful = sum(1 for r in self.test_results if r['success'])
        total = len(self.test_results)

        print(f"Total Tests: {total}")
        print(f"Successful: {successful}")
        print(f"Failed: {total - successful}")

        if successful == total:
            print("\n✅ All tests passed!")
        else:
            print("\n❌ Some tests failed. Check the details above.")

        # Check OCR specifically
        ocr_tests = [r for r in self.test_results if r.get('extraction_method') == 'ocr']
        if ocr_tests:
            print(f"\nOCR Tests: {len(ocr_tests)}")
            avg_confidence = sum(r['extraction_confidence'] for r in ocr_tests) / len(ocr_tests)
            print(f"Average OCR Confidence: {avg_confidence:.2f}")
        else:
            print("\n⚠️  No OCR tests were run. Check if OCR is properly configured.")


def main():
    parser = argparse.ArgumentParser(description='Test OMHA Exhibit Tagger OCR functionality')
    parser.add_argument('--url', default='http://localhost:8001',
                        help='Base URL of the service (default: http://localhost:8001)')
    parser.add_argument('--save-test-docs', action='store_true',
                        help='Save generated test documents to disk')

    args = parser.parse_args()

    # Run tests
    test_suite = OCRTestSuite(args.url)
    test_suite.run_all_tests()

    # Optionally save test documents
    if args.save_test_docs:
        print("\nSaving test documents to ./test_documents/")
        os.makedirs("test_documents", exist_ok=True)

        # Generate and save documents
        medicare_text = "MEDICARE RECONSIDERATION DECISION\n\nTest document for OCR testing."

        with open("test_documents/text_based.pdf", "wb") as f:
            f.write(DocumentTestGenerator.create_text_pdf(medicare_text, "text.pdf"))

        with open("test_documents/scanned.pdf", "wb") as f:
            f.write(DocumentTestGenerator.create_scanned_pdf(medicare_text, "scanned.pdf"))

        print("Test documents saved!")


if __name__ == "__main__":
    main()