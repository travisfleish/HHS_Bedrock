from typing import List, Dict, Optional, Tuple
import base64
import PyPDF2
from docx import Document
import io
import logging
import time
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from enum import Enum

# OCR imports with fallback
try:
    import pytesseract
    from pdf2image import convert_from_bytes
    from PIL import Image, ImageEnhance, ImageFilter
    import numpy as np

    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

logger = logging.getLogger(__name__)


class PDFType(Enum):
    """PDF document types"""
    TEXT_BASED = "text-based"
    SCANNED = "scanned"
    MIXED = "mixed"  # Contains both text and scanned pages


@dataclass
class ExtractionResult:
    """Result of document extraction"""
    text: str
    method: str  # 'pypdf2', 'ocr', 'mixed'
    pages_processed: int
    extraction_time: float
    confidence: float  # 0-1 confidence in extraction quality
    warnings: List[str]


class DocumentProcessor:
    """Enhanced document processor with intelligent OCR support"""

    def __init__(self,
                 ocr_enabled: bool = True,
                 ocr_dpi: int = 300,
                 min_text_threshold: int = 100,
                 parallel_ocr: bool = True,
                 max_workers: int = 4):
        """
        Initialize document processor

        Args:
            ocr_enabled: Enable OCR for scanned PDFs
            ocr_dpi: DPI for OCR conversion (higher = better quality but slower)
            min_text_threshold: Minimum characters to consider page as text-based
            parallel_ocr: Process multiple pages in parallel
            max_workers: Maximum parallel workers for OCR
        """
        self.ocr_enabled = ocr_enabled and OCR_AVAILABLE
        self.ocr_dpi = ocr_dpi
        self.min_text_threshold = min_text_threshold
        self.parallel_ocr = parallel_ocr
        self.max_workers = max_workers

        if ocr_enabled and not OCR_AVAILABLE:
            logger.warning(
                "OCR requested but dependencies not installed. Install with: pip install pytesseract pdf2image Pillow")

    def preprocess_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """Preprocess image to improve OCR accuracy"""
        try:
            # Convert to grayscale
            image = image.convert('L')

            # Enhance contrast
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(1.5)

            # Apply slight sharpening
            image = image.filter(ImageFilter.SHARPEN)

            # Convert to numpy array for advanced processing
            img_array = np.array(image)

            # Apply adaptive thresholding
            # This helps with poor quality scans and varying lighting
            threshold = np.mean(img_array)
            img_array = np.where(img_array > threshold, 255, 0).astype(np.uint8)

            # Convert back to PIL Image
            image = Image.fromarray(img_array)

            return image
        except Exception as e:
            logger.warning(f"Image preprocessing failed: {e}")
            return image

    def extract_text_from_pdf(self, content: bytes) -> Tuple[str, List[int]]:
        """
        Extract text from PDF content using PyPDF2

        Returns:
            (extracted_text, pages_with_low_text)
        """
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text_parts = []
        pages_with_low_text = []

        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            text_parts.append(page_text)

            # Track pages with minimal text (likely scanned)
            if len(page_text.strip()) < self.min_text_threshold:
                pages_with_low_text.append(page_num)

        return '\n'.join(text_parts), pages_with_low_text

    def ocr_single_page(self, image: Image.Image, page_num: int) -> str:
        """OCR a single page with preprocessing"""
        try:
            # Preprocess image
            processed_image = self.preprocess_image_for_ocr(image)

            # Perform OCR with custom config for better accuracy
            custom_config = r'--oem 3 --psm 6'  # Best OCR engine mode, uniform block of text
            text = pytesseract.image_to_string(processed_image, config=custom_config)

            logger.debug(f"OCR page {page_num}: extracted {len(text)} characters")
            return text
        except Exception as e:
            logger.error(f"OCR failed for page {page_num}: {e}")
            return ""

    def extract_text_from_pdf_with_ocr(self, content: bytes, specific_pages: Optional[List[int]] = None) -> str:
        """
        Extract text from PDF using OCR

        Args:
            content: PDF content bytes
            specific_pages: List of page numbers to OCR (0-indexed), or None for all pages
        """
        if not self.ocr_enabled:
            logger.warning("OCR is disabled or not available")
            return ""

        logger.info(f"Starting OCR extraction (DPI={self.ocr_dpi}, parallel={self.parallel_ocr})...")
        start_time = time.time()

        try:
            # Convert PDF to images
            images = convert_from_bytes(content, dpi=self.ocr_dpi)

            # Filter to specific pages if requested
            if specific_pages:
                images = [img for i, img in enumerate(images) if i in specific_pages]
                page_indices = specific_pages
            else:
                page_indices = list(range(len(images)))

            text_parts = []

            if self.parallel_ocr and len(images) > 1:
                # Parallel OCR processing
                with ThreadPoolExecutor(max_workers=min(self.max_workers, len(images))) as executor:
                    future_to_page = {
                        executor.submit(self.ocr_single_page, img, idx): (idx, img)
                        for idx, img in zip(page_indices, images)
                    }

                    # Collect results in order
                    results = {}
                    for future in as_completed(future_to_page):
                        page_idx, _ = future_to_page[future]
                        try:
                            results[page_idx] = future.result()
                        except Exception as e:
                            logger.error(f"OCR failed for page {page_idx}: {e}")
                            results[page_idx] = ""

                    # Assemble in correct order
                    for idx in sorted(results.keys()):
                        text_parts.append(results[idx])
            else:
                # Sequential OCR processing
                for idx, image in zip(page_indices, images):
                    text_parts.append(self.ocr_single_page(image, idx))

            full_text = '\n'.join(text_parts)

            elapsed_time = time.time() - start_time
            logger.info(
                f"OCR completed in {elapsed_time:.2f}s. Extracted {len(full_text)} characters from {len(images)} pages")

            return full_text

        except Exception as e:
            logger.error(f"OCR extraction failed: {str(e)}")
            return ""

    def detect_pdf_type(self, content: bytes) -> Tuple[PDFType, str, List[int]]:
        """
        Detect if PDF is text-based, scanned, or mixed

        Returns:
            (pdf_type, extracted_text, pages_needing_ocr)
        """
        text, low_text_pages = self.extract_text_from_pdf(content)

        # Analyze the extraction
        total_text_length = len(''.join(text.split()))

        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        total_pages = len(pdf_reader.pages)

        avg_text_per_page = total_text_length / total_pages if total_pages > 0 else 0

        # Classify PDF type
        if len(low_text_pages) == 0:
            logger.info(f"PDF: text-based ({total_text_length} chars, {avg_text_per_page:.0f} chars/page)")
            return PDFType.TEXT_BASED, text, []
        elif len(low_text_pages) == total_pages:
            logger.info(f"PDF: scanned (only {total_text_length} chars from {total_pages} pages)")
            return PDFType.SCANNED, "", list(range(total_pages))
        else:
            logger.info(f"PDF: mixed ({len(low_text_pages)}/{total_pages} pages need OCR)")
            return PDFType.MIXED, text, low_text_pages

    def calculate_extraction_confidence(self, text: str, method: str) -> float:
        """Calculate confidence score for extraction quality"""
        if not text:
            return 0.0

        # Basic heuristics for confidence
        confidence = 0.5  # Base confidence

        # Check for reasonable text patterns
        words = text.split()
        if len(words) > 50:
            confidence += 0.2

        # Check for sentence-like structures
        sentences = re.split(r'[.!?]+', text)
        if len(sentences) > 5:
            confidence += 0.1

        # Check for alphanumeric ratio
        alphanumeric = sum(c.isalnum() for c in text)
        if alphanumeric / len(text) > 0.7:
            confidence += 0.1

        # OCR typically has lower confidence
        if method == 'ocr':
            confidence *= 0.8

        return min(confidence, 1.0)

    def process_pdf_smart(self, content: bytes) -> ExtractionResult:
        """
        Smart PDF processing with detailed results
        """
        start_time = time.time()
        warnings = []

        pdf_type, text, pages_needing_ocr = self.detect_pdf_type(content)

        if pdf_type == PDFType.TEXT_BASED:
            method = 'pypdf2'
            pages_processed = len(PyPDF2.PdfReader(io.BytesIO(content)).pages)

        elif pdf_type == PDFType.SCANNED:
            if self.ocr_enabled:
                logger.info("Processing fully scanned PDF with OCR...")
                text = self.extract_text_from_pdf_with_ocr(content)
                method = 'ocr'
                pages_processed = len(pages_needing_ocr)
                if not text:
                    warnings.append("OCR extraction produced no text")
            else:
                warnings.append("PDF appears scanned but OCR is disabled")
                method = 'none'
                pages_processed = 0

        else:  # PDFType.MIXED
            if self.ocr_enabled:
                logger.info(f"Processing mixed PDF: OCR for pages {pages_needing_ocr}")
                ocr_text = self.extract_text_from_pdf_with_ocr(content, pages_needing_ocr)

                # Merge OCR text with existing text
                if ocr_text:
                    text = text + "\n\n" + ocr_text
                method = 'mixed'
                pages_processed = len(PyPDF2.PdfReader(io.BytesIO(content)).pages)
            else:
                warnings.append(f"{len(pages_needing_ocr)} pages appear scanned but OCR is disabled")
                method = 'pypdf2'
                pages_processed = len(PyPDF2.PdfReader(io.BytesIO(content)).pages)

        # Calculate confidence
        confidence = self.calculate_extraction_confidence(text, method)

        # Add warning for low confidence
        if confidence < 0.5:
            warnings.append("Low confidence in extraction quality")

        elapsed_time = time.time() - start_time

        return ExtractionResult(
            text=text,
            method=method,
            pages_processed=pages_processed,
            extraction_time=elapsed_time,
            confidence=confidence,
            warnings=warnings
        )

    @staticmethod
    def extract_text_from_docx(content: bytes) -> str:
        """Extract text from DOCX content"""
        doc_file = io.BytesIO(content)
        doc = Document(doc_file)

        paragraphs = [p.text for p in doc.paragraphs]

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)

        return '\n'.join(paragraphs)

    def process_document(self, content: str, mime_type: str) -> Tuple[str, ExtractionResult]:
        """
        Process document based on mime type

        Returns:
            (extracted_text, extraction_result)
        """
        # Decode base64 content
        content_bytes = base64.b64decode(content)

        # Log document info
        logger.info(f"Processing document: mime_type={mime_type}, size={len(content_bytes)} bytes")

        start_time = time.time()

        if mime_type == "application/pdf":
            # Use smart PDF processing
            result = self.process_pdf_smart(content_bytes)
            return result.text, result

        elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           "application/msword"]:
            text = self.extract_text_from_docx(content_bytes)
            result = ExtractionResult(
                text=text,
                method='python-docx',
                pages_processed=1,
                extraction_time=time.time() - start_time,
                confidence=0.95,  # DOCX extraction is usually reliable
                warnings=[]
            )
            return text, result

        elif mime_type == "text/plain":
            text = content_bytes.decode('utf-8')
            result = ExtractionResult(
                text=text,
                method='direct',
                pages_processed=1,
                extraction_time=time.time() - start_time,
                confidence=1.0,
                warnings=[]
            )
            return text, result

        else:
            # Try to decode as text
            try:
                text = content_bytes.decode('utf-8')
                result = ExtractionResult(
                    text=text,
                    method='fallback',
                    pages_processed=1,
                    extraction_time=time.time() - start_time,
                    confidence=0.7,
                    warnings=[f"Unknown mime type: {mime_type}"]
                )
                return text, result
            except Exception as e:
                logger.error(f"Failed to decode document as text: {e}")
                result = ExtractionResult(
                    text="",
                    method='failed',
                    pages_processed=0,
                    extraction_time=time.time() - start_time,
                    confidence=0.0,
                    warnings=[f"Failed to process document: {str(e)}"]
                )
                return "", result